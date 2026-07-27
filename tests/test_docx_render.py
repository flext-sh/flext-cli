"""Model-driven DOCX rendering and reading contract tests."""

from __future__ import annotations

from io import BytesIO

from docx import Document

from flext_cli import cli, m
from flext_tests import tm


def test_docx_render_is_reproducible_for_explicit_source_date_epoch() -> None:
    plan = m.Cli.DocxDocumentPlan(
        paragraphs=(m.Cli.DocxParagraphPlan(runs=(m.Cli.DocxRunPlan(text="Stable"),)),)
    )
    request = m.Cli.DocxRenderRequest(plan=plan, source_date_epoch=0)

    first = cli.docx_render(request)
    second = cli.docx_render(request)

    tm.that(first.success, eq=True, msg=first.error)
    tm.that(second.success, eq=True, msg=second.error)
    tm.that(first.value.content, eq=second.value.content)


def test_docx_render_empty_document() -> None:
    """Rendering an empty plan produces a valid DOCX."""
    plan = m.Cli.DocxDocumentPlan()
    result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(result.success, eq=True, msg=result.error)
    tm.that(result.value.content[:2], eq=b"PK")
    document = Document(BytesIO(result.value.content))
    tm.that(len(document.paragraphs), eq=0)


def test_docx_render_paragraphs() -> None:
    """Rendered paragraphs contain the supplied text."""
    plan = m.Cli.DocxDocumentPlan(
        paragraphs=(
            m.Cli.DocxParagraphPlan(runs=(m.Cli.DocxRunPlan(text="Hello"),)),
            m.Cli.DocxParagraphPlan(runs=(m.Cli.DocxRunPlan(text="World"),)),
        )
    )
    result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(result.success, eq=True, msg=result.error)
    document = Document(BytesIO(result.value.content))
    tm.that([p.text for p in document.paragraphs], eq=["Hello", "World"])


def test_docx_render_with_styles() -> None:
    """Style attributes are preserved in the rendered document."""
    plan = m.Cli.DocxDocumentPlan(
        paragraphs=(
            m.Cli.DocxParagraphPlan(
                runs=(
                    m.Cli.DocxRunPlan(
                        text="Styled",
                        style=m.Cli.DocxRunStyleSpec(
                            font=m.Cli.DocxFontSpec(
                                bold=True,
                                italic=True,
                                color=m.Cli.DocxRgbColor(value="FF0000"),
                            )
                        ),
                    ),
                ),
                alignment="center",
            ),
        )
    )
    result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(result.success, eq=True, msg=result.error)
    document = Document(BytesIO(result.value.content))
    paragraph = document.paragraphs[0]
    tm.that(paragraph.alignment, eq=1)  # WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    tm.that(run.font.bold, eq=True)
    tm.that(run.font.italic, eq=True)
    tm.that(str(run.font.color.rgb), eq="FF0000")


def test_docx_render_with_table() -> None:
    """Tables are rendered with the expected cell content."""
    plan = m.Cli.DocxDocumentPlan(
        tables=(
            m.Cli.DocxTablePlan(
                rows=(
                    m.Cli.DocxTableRowPlan(
                        cells=(
                            m.Cli.DocxTableCellPlan(
                                paragraphs=(
                                    m.Cli.DocxParagraphPlan(
                                        runs=(m.Cli.DocxRunPlan(text="A"),)
                                    ),
                                )
                            ),
                            m.Cli.DocxTableCellPlan(
                                paragraphs=(
                                    m.Cli.DocxParagraphPlan(
                                        runs=(m.Cli.DocxRunPlan(text="B"),)
                                    ),
                                )
                            ),
                        )
                    ),
                    m.Cli.DocxTableRowPlan(
                        cells=(
                            m.Cli.DocxTableCellPlan(
                                paragraphs=(
                                    m.Cli.DocxParagraphPlan(
                                        runs=(m.Cli.DocxRunPlan(text="C"),)
                                    ),
                                )
                            ),
                            m.Cli.DocxTableCellPlan(
                                paragraphs=(
                                    m.Cli.DocxParagraphPlan(
                                        runs=(m.Cli.DocxRunPlan(text="D"),)
                                    ),
                                )
                            ),
                        )
                    ),
                )
            ),
        )
    )
    result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(result.success, eq=True, msg=result.error)
    document = Document(BytesIO(result.value.content))
    tm.that(len(document.tables), eq=1)
    table = document.tables[0]
    tm.that(table.rows[0].cells[0].text, eq="A")
    tm.that(table.rows[0].cells[1].text, eq="B")
    tm.that(table.rows[1].cells[0].text, eq="C")
    tm.that(table.rows[1].cells[1].text, eq="D")


def test_docx_render_core_properties() -> None:
    """Core document properties are embedded in the rendered bytes."""
    plan = m.Cli.DocxDocumentPlan(
        paragraphs=(m.Cli.DocxParagraphPlan(runs=(m.Cli.DocxRunPlan(text="Doc"),)),),
        core_properties={"title": "Test Title", "author": "Test Author"},
    )
    result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(result.success, eq=True, msg=result.error)
    document = Document(BytesIO(result.value.content))
    tm.that(document.core_properties.title, eq="Test Title")
    tm.that(document.core_properties.author, eq="Test Author")


def test_docx_render_accounting_underline() -> None:
    """Accounting underline values are serialized and read back."""
    plan = m.Cli.DocxDocumentPlan(
        paragraphs=(
            m.Cli.DocxParagraphPlan(
                runs=(
                    m.Cli.DocxRunPlan(
                        text="Underlined",
                        style=m.Cli.DocxRunStyleSpec(
                            font=m.Cli.DocxFontSpec(underline="singleAccounting")
                        ),
                    ),
                )
            ),
        )
    )
    result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(result.success, eq=True, msg=result.error)
    read_result = cli.docx_read(result.value.content)
    tm.that(read_result.success, eq=True, msg=read_result.error)
    font = read_result.value.paragraphs[0].runs[0].style.font
    assert font is not None
    tm.that(font.underline, eq="singleAccounting")


def test_docx_read_round_trip() -> None:
    """Reading rendered bytes reproduces the document plan."""
    plan = m.Cli.DocxDocumentPlan(
        paragraphs=(
            m.Cli.DocxParagraphPlan(runs=(m.Cli.DocxRunPlan(text="Hello"),)),
            m.Cli.DocxParagraphPlan(
                runs=(m.Cli.DocxRunPlan(text="World"),), alignment="right"
            ),
        )
    )
    render_result = cli.docx_render(m.Cli.DocxRenderRequest(plan=plan))
    tm.that(render_result.success, eq=True, msg=render_result.error)
    read_result = cli.docx_read(render_result.value.content)
    tm.that(read_result.success, eq=True, msg=read_result.error)
    tm.that(read_result.value.paragraphs[0].runs[0].text, eq="Hello")
    tm.that(read_result.value.paragraphs[1].runs[0].text, eq="World")
    tm.that(read_result.value.paragraphs[1].alignment, eq="right")


def test_docx_read_rejects_invalid_bytes() -> None:
    """Reading invalid bytes returns a failure result."""
    result = cli.docx_read(b"not a docx")
    tm.that(result.success, eq=False)
    tm.that(result.error, ne=None)
